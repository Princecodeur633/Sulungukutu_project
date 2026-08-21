# -*- coding: utf-8 -*-
"""Génère le dossier de présentation Sulungukutu (DOCX + PDF)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from fpdf import FPDF

ROOT = Path(__file__).resolve().parents[2]
LOGO = ROOT / "apps" / "web" / "src" / "img" / "logo.png"
OUT_DIR = Path(__file__).resolve().parent / "livrables"
DOCX_PATH = OUT_DIR / "Sulungukutu_Dossier_Presentation_Partenaires.docx"
PDF_PATH = OUT_DIR / "Sulungukutu_Dossier_Presentation_Partenaires.pdf"

NAVY = RGBColor(0x22, 0x1F, 0x1D)
INDIGO = RGBColor(0x4F, 0x46, 0xE5)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)
EMERALD = RGBColor(0x05, 0x96, 0x69)
SKY = RGBColor(0x0E, 0xA5, 0xE9)
ROSE = RGBColor(0xDC, 0x26, 0x26)
TEXT = RGBColor(0x1C, 0x19, 0x17)
MUTED = RGBColor(0x57, 0x53, 0x4E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT = "FAFAF9"
INDIGO_HEX = "4F46E5"
NAVY_HEX = "221F1D"
AMBER_HEX = "F59E0B"
EMERALD_HEX = "059669"
ROW_ALT = "EEF2FF"
HEAD_BG = "221F1D"

FONT_REG = Path(r"C:\Windows\Fonts\calibri.ttf")
FONT_BOLD = Path(r"C:\Windows\Fonts\calibrib.ttf")
FONT_ITALIC = Path(r"C:\Windows\Fonts\calibrii.ttf")
FONT_BI = Path(r"C:\Windows\Fonts\calibriz.ttf")
FONT_TITLE = Path(r"C:\Windows\Fonts\segoeuib.ttf")
FONT_TITLE_REG = Path(r"C:\Windows\Fonts\segoeui.ttf")
FONT_TITLE_SEMI = Path(r"C:\Windows\Fonts\segoeuisb.ttf")


def shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_borders(cell, color="E7E5E4", sz="4") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def set_cell_margins(cell, **sides) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, cm_val in sides.items():
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(int(cm_val * 567)))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color=TEXT):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_text(p, text, **kwargs):
    run = p.add_run(text)
    set_run_font(run, **kwargs)
    return run


def prevent_table_break(table) -> None:
    tbl = table._tbl
    for row in tbl.tr_lst:
        tr_pr = row.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        tr_pr.append(cant)


def set_table_width(table, width_cm: float) -> None:
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_cm * 567)))
    tbl_w.set(qn("w:type"), "dxa")


def add_bottom_border(paragraph, color="4F46E5", sz="12") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)
    set_run_font(run, size=8, color=MUTED)


# ---------------------------------------------------------------------------
# Contenu
# ---------------------------------------------------------------------------

SOMMAIRE = [
    ("01", "L’essentiel"),
    ("02", "Le problème que nous résolvons"),
    ("03", "Un marché prêt, encore peu équipé"),
    ("04", "La solution Sulungukutu"),
    ("05", "L’application dans son ensemble"),
    ("06", "Ce que Sulungukutu réduit, ce qu’elle améliore"),
    ("07", "Modèle commercial et go-to-market"),
    ("08", "Pourquoi un établissement paie — et pourquoi un investisseur entre"),
    ("09", "Avantage concurrentiel"),
    ("10", "Socle technique et sécurité"),
    ("11", "État d’avancement et feuille de route"),
    ("12", "Partenariats recherchés et prochaines étapes"),
]


def build_docx() -> None:
    doc = Document()

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.9)
        section.right_margin = Cm(1.9)
        section.header_distance = Cm(0.8)
        section.footer_distance = Cm(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = TEXT
    styles["Normal"].paragraph_format.space_after = Pt(8)
    styles["Normal"].paragraph_format.line_spacing = 1.15

    # ---- Cover (section 0, no header) ----
    cover_section = doc.sections[0]
    cover_section.top_margin = Cm(0)
    cover_section.bottom_margin = Cm(0)
    cover_section.left_margin = Cm(0)
    cover_section.right_margin = Cm(0)
    cover_section.header.is_linked_to_previous = False
    cover_section.footer.is_linked_to_previous = False
    cover_section.header.paragraphs[0].text = ""
    cover_section.footer.paragraphs[0].text = ""

    cover = doc.add_table(rows=1, cols=1)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(cover, 21.0)
    cell = cover.cell(0, 0)
    shade_cell(cell, NAVY_HEX)
    set_cell_margins(cell, top=1.8, bottom=1.6, left=2.0, right=1.8)
    set_cell_borders(cell, color=NAVY_HEX, sz="0")
    tr_pr = cover.rows[0]._tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), "16838")
    tr_height.set(qn("w:hRule"), "exact")
    tr_pr.append(tr_height)
    v_align = OxmlElement("w:vAlign")
    v_align.set(qn("w:val"), "top")
    cell._tc.get_or_add_tcPr().append(v_align)

    def cp(text, size=11, bold=False, color=WHITE, space_after=6, space_before=0, align="left", name="Calibri"):
        p = cell.add_paragraph()
        p.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }[align]
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.line_spacing = 1.12
        add_text(p, text, name=name, size=size, bold=bold, color=color)
        return p

    if LOGO.exists():
        p_logo = cell.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_logo.paragraph_format.space_after = Pt(10)
        run = p_logo.add_run()
        run.add_picture(str(LOGO), width=Cm(2.0))
    else:
        cell.paragraphs[0].text = ""

    cp("DOCUMENT CONFIDENTIEL  ·  DESTINÉ AUX PARTENAIRES ET INVESTISSEURS", 9, True, AMBER, 18, 8)
    cp("SULUNGUKUTU", 32, True, WHITE, 6, 4, name="Calibri")
    cp("Dossier de présentation", 18, False, RGBColor(0xC7, 0xD2, 0xFE), 10)
    amber_line = cell.add_paragraph()
    amber_line.paragraph_format.space_after = Pt(14)
    add_text(amber_line, "━━━━━━━━━━━━━━━━━━━━", size=10, color=AMBER)
    cp("La plateforme de gestion scolaire conçue pour les établissements d’Afrique centrale.", 13, False, WHITE, 8)
    cp("Un seul outil pour piloter la pédagogie, la scolarité et la relation avec les familles — et rendre enfin visible l’argent, les absences et les résultats.", 12, False, RGBColor(0xD6, 0xD3, 0xD1), 18)

    for label, value in [
        ("Promoteur", "Kassy Gloire Exaucé"),
        ("Produit", "Plateforme SaaS multi-établissements — version 1.0 opérationnelle"),
        ("Marché prioritaire", "Écoles privées du Congo-Brazzaville, puis CEMAC"),
        ("Document", "Dossier partenaires  ·  août 2026"),
    ]:
        row_p = cell.add_paragraph()
        row_p.paragraph_format.space_after = Pt(4)
        add_text(row_p, f"{label}   ", size=10, bold=True, color=AMBER)
        add_text(row_p, value, size=10, color=WHITE)

    cp("« Le bulletin arrive avant la fin du trimestre. »", 12, True, RGBColor(0xC7, 0xD2, 0xFE), 0, 28)

    # New section for body
    new_section = doc.add_section()
    new_section.page_width = Cm(21.0)
    new_section.page_height = Cm(29.7)
    new_section.top_margin = Cm(2.2)
    new_section.bottom_margin = Cm(2.0)
    new_section.left_margin = Cm(1.9)
    new_section.right_margin = Cm(1.9)
    new_section.header_distance = Cm(0.7)
    new_section.footer_distance = Cm(0.6)
    new_section.header.is_linked_to_previous = False
    new_section.footer.is_linked_to_previous = False

    header = new_section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(4)
    add_text(hp, "SULUNGUKUTU", size=9, bold=True, color=INDIGO)
    add_text(hp, "   ·   Dossier de présentation partenaires", size=9, color=MUTED)
    add_bottom_border(hp, "4F46E5", "12")

    footer = new_section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(fp, "Confidentiel  ·  Ne pas diffuser sans autorisation  ·  Page ", size=8, color=MUTED)
    add_page_number(fp)

    usable = 17.2

    def h1(text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        add_text(p, text, size=16, bold=True, color=INDIGO)
        add_bottom_border(p, "E0E7FF", "18")

    def h2(text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        add_text(p, text, size=12.5, bold=True, color=NAVY)

    def para(text: str, italic=False, bold=False, size=11, color=TEXT, after=8) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        add_text(p, text, size=size, bold=bold, italic=italic, color=color)

    def bullets(items: list[str]) -> None:
        for item in items:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.12
            add_text(p, "▸  ", size=11, bold=True, color=INDIGO)
            add_text(p, item, size=11, color=TEXT)

    def callout(title: str, body: str, fill="EEF2FF", accent=INDIGO_HEX) -> None:
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_width(table, usable)
        c = table.cell(0, 0)
        shade_cell(c, fill)
        set_cell_margins(c, top=0.25, bottom=0.25, left=0.35, right=0.35)
        set_cell_borders(c, color=accent, sz="8")
        p1 = c.paragraphs[0]
        p1.paragraph_format.space_after = Pt(4)
        add_text(p1, title, size=10, bold=True, color=INDIGO)
        p2 = c.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.12
        add_text(p2, body, size=10.5, color=TEXT)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def add_table(headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None) -> None:
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_width(table, usable)
        prevent_table_break(table)
        if col_widths:
            for i, w in enumerate(col_widths):
                for cell in table.columns[i].cells:
                    cell.width = Cm(w)
        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            shade_cell(cell, HEAD_BG)
            set_cell_margins(cell, top=0.12, bottom=0.12, left=0.18, right=0.18)
            set_cell_borders(cell, color=HEAD_BG, sz="4")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_text(p, h, size=9, bold=True, color=WHITE)
        for r_i, row in enumerate(rows):
            for c_i, val in enumerate(row):
                cell = table.cell(r_i + 1, c_i)
                shade_cell(cell, ROW_ALT if r_i % 2 == 0 else "FFFFFF")
                set_cell_margins(cell, top=0.1, bottom=0.1, left=0.18, right=0.18)
                set_cell_borders(cell, color="E7E5E4", sz="4")
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.08
                add_text(p, val, size=9.5, color=TEXT)
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(8)

    def kpi_row(items: list[tuple[str, str, str]]) -> None:
        table = doc.add_table(rows=1, cols=len(items))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_width(table, usable)
        w = usable / len(items)
        for i, (k, v, sub) in enumerate(items):
            cell = table.cell(0, i)
            cell.width = Cm(w)
            shade_cell(cell, "EEF2FF" if i % 2 == 0 else "FFFBEB")
            set_cell_margins(cell, top=0.22, bottom=0.22, left=0.22, right=0.22)
            set_cell_borders(cell, color="E0E7FF", sz="6")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            add_text(p, k.upper(), size=8, bold=True, color=INDIGO)
            p2 = cell.add_paragraph()
            p2.paragraph_format.space_after = Pt(2)
            add_text(p2, v, size=14, bold=True, color=NAVY)
            p3 = cell.add_paragraph()
            p3.paragraph_format.space_after = Pt(0)
            add_text(p3, sub, size=8, color=MUTED)
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ---- Sommaire ----
    h1("Sommaire")
    para("Ce document présente Sulungukutu dans son ensemble : le problème, le produit, le modèle économique, l’impact opérationnel et les conditions d’un partenariat ou d’un investissement.")
    toc = doc.add_table(rows=len(SOMMAIRE), cols=2)
    set_table_width(toc, usable)
    toc.columns[0].width = Cm(1.6)
    toc.columns[1].width = Cm(15.6)
    for i, (num, title) in enumerate(SOMMAIRE):
        c0, c1 = toc.cell(i, 0), toc.cell(i, 1)
        shade_cell(c0, "FFFFFF")
        shade_cell(c1, "FFFFFF")
        set_cell_borders(c0, "FFFFFF", "0")
        set_cell_borders(c1, "FFFFFF", "0")
        set_cell_margins(c0, top=0.05, bottom=0.05, left=0.05, right=0.1)
        set_cell_margins(c1, top=0.05, bottom=0.05, left=0.05, right=0.05)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        add_text(p0, num, size=11, bold=True, color=INDIGO)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        add_text(p1, title, size=11, color=TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---- 01 ----
    h1("01  ·  L’essentiel")
    para("Sulungukutu est une plateforme numérique de gestion scolaire, pensée pour les établissements d’Afrique subsaharienne et d’abord calibrée sur le système éducatif congolais. Elle remplace le trio encore dominant — cahiers, fichiers Excel, groupes WhatsApp — par une source unique de vérité, partagée entre la direction, les enseignants, les parents et les élèves.")
    para("Le produit n’est pas une intention. La version 1.0 est développée, déployable, et couvre déjà le cycle de vie d’une année scolaire : inscriptions, classes, notes, présences, bulletins trimestriels, mensualités, messagerie, annonces, exports et journal d’audit. L’architecture est multi-établissements : un promoteur peut piloter plusieurs écoles depuis un même compte.")
    kpi_row([
        ("Produit", "Version 1.0", "Opérationnelle, prête pour pilotes"),
        ("Utilisateurs", "5 espaces", "Direction, enseignants, parents, élèves, super-admin"),
        ("Cœur métier", "Pédagogie + scolarité", "Notes, bulletins, présences, paiements"),
        ("Ambition", "SaaS régional", "Congo d’abord, CEMAC ensuite"),
    ])
    callout(
        "La promesse, en une phrase",
        "Donner à chaque établissement une vision en temps réel de ses élèves, de ses résultats et de sa trésorerie — et redonner aux familles le droit de savoir, sans attendre le conseil de classe.",
    )
    para("Nous présentons ce dossier à des collaborateurs, des partenaires stratégiques et des investisseurs potentiels. L’objectif n’est pas de vendre un rêve : c’est de montrer un outil déjà construit, un marché réel, un modèle commercial simple, et une fenêtre d’entrée encore ouverte.")

    # ---- 02 ----
    h1("02  ·  Le problème que nous résolvons")
    h2("Une école privée congolaise gagne ou perd sur la gestion, pas seulement sur la pédagogie")
    para("Au Congo-Brazzaville, le privé n’est plus un complément. Selon les données publiques récentes, environ 41 % des élèves de l’enseignement général fréquentent un établissement privé (PASEC / MEPPSA). En primaire, près d’un établissement sur deux est privé, et près de la moitié des élèves y sont scolarisés. En 2026, 1 192 établissements privés d’enseignement général étaient soumis à la commission d’agrément.")
    para("Ces écoles vivent pourtant encore, pour la plupart, dans un régime artisanal :")
    bullets([
        "Les notes circulent dans des cahiers, puis sont recopiées, puis recalculées à la main au moment des bulletins.",
        "Les présences se marquent le matin et s’évaporent le soir : le parent apprend une absence trop tard, parfois jamais.",
        "La scolarité se suit sur un registre ou un Excel. Les acomptes, les exonérations et les relances se perdent. L’impayé n’est pas un incident : c’est un trou de trésorerie structurel.",
        "La communication passe par WhatsApp, sans traçabilité, sans ciblage, sans preuve.",
        "Le promoteur pilote à l’intuition. Il découvre les décrochages, les classes en difficulté et les retards de paiement quand il est trop tard pour agir.",
    ])
    para("Ce n’est pas un détail opérationnel. C’est le cœur du risque économique de l’école privée. Les institutions financières locales (CAPPED, HOPE Congo et d’autres) proposent précisément des crédits aux promoteurs « pour payer les salaires sans attendre que tous les parents aient réglé ». Autrement dit : le recouvrement est le talon d’Achille du secteur — et personne n’a encore industrialisé sa visibilité.")
    add_table(
        ["Acteur", "Ce qu’il vit aujourd’hui", "Ce qu’il devrait vivre"],
        [
            ["Promoteur / Directeur", "Pilotage à l’aveugle, trésorerie stressée, image artisanale", "Tableau de bord du jour : effectifs, absences, impayés, moyennes"],
            ["Secrétariat", "Saisies multiples, relances papier, bulletins de dernière minute", "Une base unique, des exports, des bulletins générés"],
            ["Enseignant", "Cahier de notes, appels répétitifs, peu de feedback", "Saisie simple, présences, classes sous les yeux"],
            ["Parent", "Découvre les résultats trop tard, doute sur les paiements", "Suivi de l’enfant, bulletin publié, historique de scolarité"],
            ["Élève", "Peu de visibilité sur son propre parcours", "Notes, présences, bulletin, emploi du temps"],
        ],
        [3.4, 6.9, 6.9],
    )
    callout(
        "Le coût invisible",
        "Chaque heure perdue au secrétariat, chaque mensualité « oubliée », chaque bulletin publié en retard, chaque parent qui retire son enfant faute de confiance : ce n’est pas de l’intendance. C’est du chiffre d’affaires, de la réputation et de la qualité pédagogique qui s’évaporent.",
        fill="FFFBEB",
        accent=AMBER_HEX,
    )

    # ---- 03 ----
    h1("03  ·  Un marché prêt, encore peu équipé")
    h2("Un socle national concentré et adressable")
    para("Le marché congolais a une caractéristique rare pour un investisseur : il est à la fois suffisamment grand pour construire une entreprise, et suffisamment concentré pour être attaqué avec une équipe courte.")
    add_table(
        ["Signal de marché", "Donnée", "Implication pour Sulungukutu"],
        [
            ["Écoles privées d’enseignement général", "1 192 dossiers d’agrément (2026)", "Cible B2B identifiable, organisée en associations"],
            ["Poids du privé", "≈ 41 % des élèves (général) ; ≈ 49 % en primaire", "Le privé est central, pas périphérique"],
            ["Géographie", "89 % des écoles privées à Brazzaville et Pointe-Noire", "Deux villes suffisent pour un premier réseau dense"],
            ["Ratio d’encadrement privé", "≈ 27 élèves / enseignant (vs 48 dans le public)", "Des écoles déjà plus « gérables » numériquement"],
            ["Pression réglementaire", "Agrément, mise en conformité, modernisation", "La traçabilité devient un argument, pas un luxe"],
            ["Digitalisation en cours", "Sessions promoteurs, outils locaux encore émergents", "La question n’est plus « faut-il digitaliser ? » mais « avec qui ? »"],
        ],
        [4.2, 5.8, 7.2],
    )
    h2("Une fenêtre, pas un océan rouge")
    para("Des solutions existent — registres papier, Excel, quelques logiciels locaux présentés aux promoteurs, et des suites étrangères (Pronote, systèmes européens) inadaptées au Mobile Money, aux mensualités de septembre à mai, et au référentiel national. Le marché n’est pas vide. Il est fragmenté, peu saturé, et en demande de recouvrement autant que de pédagogie.")
    para("La CEMAC offre ensuite une expansion naturelle : mêmes langues de travail, même culture des trimestres et des coefficients, même pénétration du Mobile Money, mêmes promoteurs multi-sites. Le Congo n’est pas le plafond. C’est la tête de pont.")
    callout(
        "Lecture investisseur",
        "Nous ne promettons pas un « marché de plusieurs milliards » pour impressionner. Nous montrons un premier pays où plus de mille écoles privées existent, où la moitié des élèves du primaire sont déjà dans le privé, où deux villes concentrent la demande, et où le paiement des familles est le nerf de la guerre. C’est un marché vendable, pas un slogan.",
    )

    # ---- 04 ----
    h1("04  ·  La solution Sulungukutu")
    para("Sulungukutu est un logiciel en ligne (SaaS), accessible depuis un navigateur, sans installation lourde dans l’école. Chaque établissement dispose de son espace isolé, de son identité visuelle (logo, couleur), de son année scolaire, et de ses utilisateurs. Un même adulte — promoteur, enseignant, parent — peut appartenir à plusieurs écoles et basculer d’un établissement à l’autre.")
    h2("Ce qui la rend différente d’un « logiciel scolaire » générique")
    bullets([
        "Elle est née du terrain congolais : trimestres T1 / T2 / T3, mentions, coefficients, séries du lycée (A, C, D, TI), mensualités de septembre à mai, paiements espèces, Mobile Money, virement, chèque, exonérations et acomptes.",
        "Le référentiel pédagogique national (CP1 → Terminale) est intégré. L’école n’a pas à recréer le programme : elle active les niveaux et les matières dont elle a besoin, puis ajuste coefficients et volumes horaires localement.",
        "Elle relie pédagogie et argent. Un bulletin, une absence, une mensualité impayée ne vivent plus dans trois univers séparés.",
        "Elle parle aux parents, pas seulement au secrétariat. Le parent voit les résultats publiés, les présences, l’historique de scolarité, et peut échanger avec l’école.",
        "Elle est conçue pour des utilisateurs sans e-mail personnel : connexion par identifiant, e-mail ou téléphone ; réinitialisation du mot de passe en présentiel par l’administration — un détail décisif en Afrique, souvent oublié par les outils importés.",
    ])
    h2("Cinq espaces, une seule base")
    add_table(
        ["Espace", "Pour qui", "Ce qu’il permet"],
        [
            ["Administration", "Directeur, promoteur, secrétariat", "Pilotage, élèves, enseignants, notes, bulletins, paiements, emplois du temps, annonces, exports, audit"],
            ["Enseignant", "Professeurs", "Notes (devoir, contrôle, examen, interro), présences, bulletins, emploi du temps, messagerie"],
            ["Parent", "Pères, mères, tuteurs", "Suivi des enfants, bulletins publiés, scolarité, présences, messages"],
            ["Élève", "Collégiens et lycéens", "Notes, bulletin, présences, emploi du temps, messagerie"],
            ["Super-admin", "Opérateur Sulungukutu / groupe scolaire", "Toutes les écoles, accès, statistiques consolidées"],
        ],
        [3.2, 4.4, 9.6],
    )

    # ---- 05 ----
    h1("05  ·  L’application dans son ensemble")
    para("Sulungukutu n’est pas un module isolé. C’est une suite complète, déjà assemblée, qui couvre le quotidien d’un établissement du premier jour de rentrée jusqu’à la publication des bulletins.")

    h2("Pilotage")
    para("Le tableau de bord direction affiche en temps réel les effectifs, le taux de présence du jour, les impayés du mois, les moyennes par classe, la répartition des mentions et l’activité récente. Des alertes intelligentes et un onboarding guidé accompagnent les nouvelles écoles. Le promoteur cesse de « demander au secrétariat » : il voit.")

    h2("Organisation pédagogique")
    bullets([
        "Niveaux, classes, matières, coefficients, volumes horaires.",
        "Activation du référentiel national (primaire, collège, lycée) sans ressaisie.",
        "Emplois du temps, avec détection des conflits d’un enseignant avant création d’un créneau.",
        "Calendrier de l’établissement.",
        "Transfert d’élève d’une classe à une autre en cours d’année.",
    ])

    h2("Élèves et familles")
    bullets([
        "Inscription d’un élève avec création automatique du compte parent.",
        "Lien père / mère / tuteur, y compris un parent déjà connu pour un autre enfant.",
        "Import CSV pour les rentrées (plus de ressaisie de listes).",
        "Annuaire des identifiants, invitations enseignants par e-mail, gestion des statuts (actif, inactif, suspendu).",
        "Identité visuelle de l’école : nom, logo, couleur d’accent.",
    ])

    h2("Notes, présences, bulletins")
    bullets([
        "Saisie unitaire ou en masse : devoir, contrôle, examen, interrogation.",
        "Appel par matière et par jour : présent, absent, retard — avec statistiques.",
        "Génération des bulletins trimestriels, calcul des moyennes et mentions, publication contrôlée, archivage, PDF.",
        "Le parent et l’élève ne voient le bulletin que lorsqu’il est publié : la direction garde la main.",
    ])

    h2("Scolarité et trésorerie")
    bullets([
        "Neuf mensualités par année scolaire, statut payé / partiel / impayé / exonéré / annulé.",
        "Journal de transactions : espèces au guichet, Mobile Money (y compris parcours à distance), virement, chèque.",
        "Acomptes, annulation d’une transaction, historique consultable.",
        "Liste des élèves impayés du mois, relances e-mail automatisables.",
        "Export Excel du recouvrement — le document que le promoteur emporte à la banque.",
    ])

    h2("Communication")
    bullets([
        "Messagerie interne entre administration, enseignants, parents et élèves.",
        "Annonces ciblées : toute l’école, parents, enseignants, élèves.",
        "Notifications (paiement, absence, bulletin, message, système) avec temps réel.",
    ])

    h2("Preuve, conformité, industrialisation")
    bullets([
        "Journal d’audit : qui a modifié un paiement, publié un bulletin, invité un utilisateur, saisi une note.",
        "Exports Excel (élèves, notes, paiements).",
        "Isolation stricte entre écoles : une école ne voit jamais les données d’une autre.",
        "Rôles et permissions, sessions révocables, limitation des tentatives de connexion.",
    ])

    # ---- 06 ----
    h1("06  ·  Ce que Sulungukutu réduit, ce qu’elle améliore")
    h2("Elle réduit")
    add_table(
        ["Elle réduit…", "Comment", "Effet pour l’école"],
        [
            ["Le temps administratif", "Bulletins générés, listes exportées, relances outillées, import de rentrée", "Le secrétariat cesse d’être un goulot d’étranglement"],
            ["Les impayés « invisibles »", "Statut de chaque mensualité, liste des retardataires, historique des acomptes", "La trésorerie redevient lisible, donc actionnable"],
            ["Les erreurs de moyennes", "Calcul automatisé, mentions, publication contrôlée", "Moins de contestations, plus de crédibilité"],
            ["Le papier et les allers-retours", "Espaces parent / élève, PDF, messagerie, annonces", "Moins d’impressions, moins de files au secrétariat"],
            ["La perte d’information", "Une base unique à la place de cahiers et d’Excel dispersés", "Plus de « le fichier est chez Madame X »"],
            ["Le délai parent–école", "Présences, notes publiées, notifications", "La confiance se construit en continu, pas au trimestre"],
            ["Le risque de contestation", "Journal d’audit, transactions horodatées, rôles", "Chaque acte sensible laisse une trace"],
            ["Le travail en double", "Un élève, un parent, une note, un paiement : saisis une fois", "Moins de fatigue, moins d’incohérences"],
        ],
        [4.0, 6.8, 6.4],
    )
    h2("Elle améliore")
    add_table(
        ["Elle améliore…", "Pour qui", "Pourquoi c’est décisif"],
        [
            ["Le recouvrement de la scolarité", "Promoteur", "C’est le salaire des enseignants et la survie de l’école"],
            ["La transparence", "Parents", "Un parent informé reste, recommande, paie mieux"],
            ["Le suivi pédagogique", "Enseignants & direction", "On voit les classes qui décrochent avant l’examen"],
            ["La qualité des décisions", "Direction", "KPI du jour plutôt que réunion de crise"],
            ["L’image de l’établissement", "Tout le monde", "Une école digitale inspire confiance aux familles urbaines"],
            ["La fidélisation des élèves", "Promoteur", "Partir coûte plus cher qu’un abonnement logiciel"],
            ["Les conditions de travail", "Personnel", "Moins de paperasse, plus de métier"],
            ["La conformité et l’agrément", "Promoteur", "Traçabilité des effectifs, des actes et des paiements"],
        ],
        [4.6, 3.4, 9.2],
    )

    h2("Scénario illustratif — collège privé de 320 élèves à Brazzaville")
    para("Les chiffres ci-dessous sont un scénario de travail, pas une garantie de résultat. Ils servent à rendre concret le retour sur investissement pour un promoteur — et, par ricochet, la raison pour laquelle l’école paiera l’abonnement.", italic=True, size=10, color=MUTED)
    add_table(
        ["Hypothèse", "Ordre de grandeur"],
        [
            ["Effectif", "320 élèves"],
            ["Mensualité moyenne", "22 000 FCFA"],
            ["Recette théorique mensuelle", "7,04 millions FCFA"],
            ["Retards / impayés chroniques (hypothèse 18 %)", "≈ 1,27 million FCFA « gelés » chaque mois"],
            ["Si Sulungukutu ne fait gagner que 6 points de recouvrement", "≈ 422 000 FCFA / mois  ·  ≈ 3,8 millions FCFA / an"],
            ["Coût logiciel (offre Professionnel, indicatif)", "45 000 FCFA / mois  ·  540 000 FCFA / an"],
            ["Ordre de grandeur du multiple", "Le seul recouvrement peut couvrir le logiciel plusieurs fois — avant même le temps gagné au secrétariat"],
        ],
        [9.0, 8.2],
    )
    callout(
        "Ce que l’investisseur doit retenir",
        "Sulungukutu ne se vend pas comme un « plus pédagogique ». Elle se vend comme un instrument de trésorerie, de réputation et de temps. C’est pour cela qu’un promoteur peut justifier la dépense, et qu’un partenaire peut y voir un revenu récurrent défendable.",
        fill="ECFDF5",
        accent=EMERALD_HEX,
    )

    # ---- 07 ----
    h1("07  ·  Modèle commercial et go-to-market")
    h2("Comment Sulungukutu gagne de l’argent")
    para("Le modèle est un abonnement SaaS facturé à l’établissement, selon la taille. Simple à expliquer, récurrent, prévisible. L’école n’achète pas une licence figée : elle paie un service vivant (mises à jour, support, hébergement, nouvelles fonctions).")
    add_table(
        ["Offre", "Cible", "Tarif indicatif", "Inclus"],
        [
            ["Essentiel", "Écoles jusqu’à 200 élèves", "20 000 FCFA / mois", "Tous les modules cœur, 1 établissement, support e-mail"],
            ["Professionnel", "200 à 600 élèves", "45 000 FCFA / mois", "Tout Essentiel + imports, exports avancés, onboarding assisté"],
            ["Groupe scolaire", "Plusieurs sites / + 600 élèves", "Sur devis + 80 FCFA / élève / mois", "Super-admin, consolidé, accompagnement dédié"],
            ["Mise en service", "Toutes les offres", "50 000 FCFA (une fois)", "Paramétrage, import rentrée, formation secrétariat + direction"],
        ],
        [3.2, 4.0, 3.8, 6.2],
    )
    para("Remise de 15 % en paiement annuel anticipé — aligné sur le rythme de trésorerie de la rentrée. Les tarifs ci-dessus sont une grille de travail, destinée à ancrer la discussion : ils restent ajustables selon le partenaire, le volume et le niveau d’accompagnement.")
    h2("Revenus additionnels, dans un second temps")
    bullets([
        "Commission discrète sur les paiements Mobile Money en production (quelques points, jamais au détriment de l’école).",
        "SMS de relance et d’absence (facturés au volume, ou inclus dans une offre supérieure).",
        "Modules avancés : paie du personnel, cantine, transport, bibliothèque.",
        "Offre « réseau » pour associations de promoteurs ou groupes confessionnels.",
    ])
    h2("Comment nous commercialisons")
    para("Le cycle de vente scolaire est saisonnier. On signe entre mai et septembre. On enchante à la rentrée. On renouvelle l’année suivante si les bulletins du T1 sont sortis sans douleur et si le recouvrement de novembre est lisible. Notre go-to-market suit cette réalité.")
    add_table(
        ["Phase", "Horizon", "Objectif", "Moyens"],
        [
            ["Pilotes", "6 à 9 mois", "10 à 15 écoles à Brazzaville et Pointe-Noire", "Démo live, tarif pilote, accompagnement intensif, preuves chiffrées"],
            ["Réseau", "12 à 24 mois", "50+ établissements au Congo", "CONAPEPCO / UCOPEP / FISEP, promoteurs multi-sites, bouche-à-oreille"],
            ["Partenaires de distribution", "En parallèle", "Canal telco & financier", "Airtel / MTN (Mobile Money), microfinance scolaire, inspecteurs et agréments"],
            ["Expansion CEMAC", "24 à 36 mois", "Gabon, RDC, Cameroun", "Même produit, référentiels nationaux additionnels, relais locaux"],
        ],
        [3.2, 2.8, 5.0, 6.2],
    )
    h2("Ordre de grandeur de revenus (scénario prudent)")
    para("Ces projections sont indicatives. Elles n’intègrent pas encore la commission Mobile Money. Elles illustrent le levier d’un modèle récurrent, même avec un rythme commercial réaliste.", italic=True, size=10, color=MUTED)
    add_table(
        ["Horizon", "Écoles actives", "Panier moyen mensuel", "ARR indicatif"],
        [
            ["Année 1", "25", "35 000 FCFA", "≈ 10,5 millions FCFA"],
            ["Année 2", "70", "38 000 FCFA", "≈ 31,9 millions FCFA"],
            ["Année 3", "140", "42 000 FCFA", "≈ 70,6 millions FCFA"],
        ],
        [3.4, 3.6, 5.0, 5.2],
    )
    para("À 140 écoles, Sulungukutu n’a encore adressé qu’environ 12 % du seul vivier privé congolais d’enseignement général. Le plafond local n’est pas atteint. L’expansion régionale n’est même pas nécessaire pour justifier une première brique d’entreprise — elle en accélère simplement l’ampleur.")

    # ---- 08 ----
    h1("08  ·  Pourquoi un établissement paie — et pourquoi un investisseur entre")
    h2("Le motif d’achat de l’école")
    para("Un promoteur n’achète pas « du digital ». Il achète trois choses qu’il peut défendre devant son conseil, son épouse associée, ou son banquier :")
    bullets([
        "Voir l’argent : qui a payé, qui n’a pas payé, depuis quand, pour quel enfant.",
        "Tenir la promesse aux parents : un bulletin à l’heure, une absence signalée, une école qui a l’air tenue.",
        "Diminuer la charge du secrétariat sans embaucher une personne de plus.",
    ])
    h2("Le motif d’entrée de l’investisseur ou du partenaire")
    bullets([
        "Le risque produit est déjà largement derrière nous : la plateforme existe, les cinq rôles existent, le métier école est modélisé.",
        "Le revenu est récurrent, collant (on ne change pas de logiciel en janvier), et saisonnier de façon prévisible.",
        "Le paiement de la scolarité est un flux réel, mensuel, massif. S’y brancher (visibilité, puis Mobile Money) crée un actif difficile à déloger.",
        "Le marché est localement concentré : on peut gagner une ville, puis un pays, sans brûler du capital en acquisition digitale mondiale.",
        "L’impact est lisible : des écoles mieux gérées, des parents mieux informés, des enseignants moins submergés. C’est un dossier qui se raconte aussi bien à un fonds qu’à un opérateur telecom ou à un réseau de promoteurs.",
        "La barrière se construit avec le temps : données de l’année scolaire, habitudes des enseignants, comptes parents, historique de paiements. Plus l’école reste, plus elle reste.",
    ])
    callout(
        "Ce n’est pas un projet « à construire si vous financez »",
        "Trop de dossiers africains demandent de l’argent pour commencer à coder. Ici, l’argent — et les alliances — servent à vendre, à accompagner, à connecter le Mobile Money réel, et à transformer un produit abouti en standard de fait des écoles privées congolaises.",
    )

    # ---- 09 ----
    h1("09  ·  Avantage concurrentiel")
    add_table(
        ["Alternative", "Limite", "Ce que Sulungukutu oppose"],
        [
            ["Cahiers + Excel + WhatsApp", "Fragile, non partagé, non auditable, lent", "Une source de vérité, des rôles, un historique"],
            ["Logiciel étranger (Pronote, etc.)", "Cher, mal calé sur le programme local, peu Mobile Money", "Référentiel congolais, mensualités, prix africain"],
            ["Outil local partiel", "Souvent un module (caisse ou notes), peu d’espace parent moderne", "Suite complète, cinq espaces, multi-écoles"],
            ["Développement interne de l’école", "Coûteux, jamais fini, dépendant d’une personne", "Produit maintenu, sécurisé, déjà prêt"],
        ],
        [4.2, 6.4, 6.6],
    )
    h2("Moats que nous renforçons volontairement")
    bullets([
        "Référentiel national intégré — plus on l’enrichit (examens d’État, grilles officielles), plus un concurrent générique est hors-jeu.",
        "Multi-établissements natif — les promoteurs sérieux ont plusieurs sites ; un outil « une école = une base » les perd.",
        "Couche parents — c’est elle qui rend le logiciel collant, pas seulement le secrétariat.",
        "Journal d’audit et isolation des données — arguments de confiance pour les groupes et, demain, pour le régulateur.",
        "Distribution par les réseaux de promoteurs — le commercial africain de l’éducation se gagne en salle, pas en bannière.",
    ])

    # ---- 10 ----
    h1("10  ·  Socle technique et sécurité")
    para("Un partenaire technologique ou un investisseur due-diligence doit pouvoir vérifier que Sulungukutu n’est pas un prototype fragile. L’architecture est celle d’un produit moderne, déjà pensé pour plusieurs écoles et pour la production.")
    add_table(
        ["Couche", "Choix", "Pourquoi c’est un atout"],
        [
            ["Application", "Next.js 14, interface soignée, cinq espaces", "Expérience contemporaine, mobile-friendly"],
            ["API", "GraphQL, temps réel (notifications, messages, paiements)", "Un écran direction qui se met à jour sans recharger"],
            ["Données", "PostgreSQL, isolation par établissement", "Sérieux, portable, auditable"],
            ["Identité", "JWT, rôles, révocation après changement de mot de passe", "Sessions maîtrisées"],
            ["Métiers sensibles", "PDF bulletins, exports Excel, imports CSV, e-mails", "Le quotidien réel de l’école, pas une démo"],
            ["Déploiement", "Cloud (web + API + base managée)", "Mise en production sans parc serveur dans chaque école"],
        ],
        [3.4, 6.4, 7.4],
    )
    bullets([
        "Isolation multi-école : une action ne peut pas modifier « l’autre établissement » par erreur de paramètre.",
        "Contrôle d’accès sur les fichiers (PDF, exports) : pas de lien public oublié.",
        "Limitation des tentatives de connexion.",
        "Mot de passe oublié par lien signé, ou reset en présentiel pour les comptes sans e-mail réel.",
        "Journal d’audit des actes sensibles (paiements, bulletins, utilisateurs).",
    ])
    para("En langage clair : Sulungukutu est déjà au niveau où l’on peut ouvrir une école pilote sans reconstruire le moteur. Le travail restant est de commercialisation, d’intégrations de paiement réelles, et d’accompagnement — pas de « refaire le logiciel ».")

    # ---- 11 ----
    h1("11  ·  État d’avancement et feuille de route")
    h2("Ce qui est déjà là — sans maquillage")
    para("Nous assumons le stade : produit abouti, commercialisation à industrialiser. C’est un stade sain. Il évite de lever de l’argent pour découvrir, deux ans plus tard, que le métier école est plus complexe qu’un tableau Kanban.")
    add_table(
        ["Brique", "Statut"],
        [
            ["Espaces direction, enseignant, parent, élève, super-admin", "Livré"],
            ["Référentiel national Congo (niveaux, séries, matières, coefficients)", "Livré"],
            ["Notes, présences, bulletins PDF, publication", "Livré"],
            ["Mensualités, acomptes, modes de règlement, parcours Mobile Money (socle)", "Livré"],
            ["Messagerie, annonces, notifications temps réel", "Livré"],
            ["Onboarding école, imports, exports, audit", "Livré"],
            ["Identité, sécurité, multi-établissements", "Livré"],
            ["Premiers établissements payants en production", "À enclencher avec les partenaires de ce dossier"],
            ["Mobile Money opérateur (Airtel Money / MTN MoMo) en réel", "Prochaine étape prioritaire"],
            ["Application mobile native parents / enseignants", "Feuille de route"],
            ["SMS de masse", "Feuille de route"],
        ],
        [11.2, 6.0],
    )
    h2("Feuille de route")
    add_table(
        ["Période", "Priorité", "Résultat visé"],
        [
            ["0 – 6 mois", "Pilotes Brazzaville & Pointe-Noire, formation, preuves d’usage", "Études de cas : temps gagné, recouvrement, satisfaction parents"],
            ["6 – 12 mois", "Mobile Money production, SMS relances, offre annuelle rentrée", "Paiement à distance réel + renouvellements"],
            ["12 – 24 mois", "Réseau 50+ écoles, app mobile parents, second pays CEMAC", "Standard de fait au Congo privé urbain"],
            ["24 – 36 mois", "Modules RH / cantine, analytique avancée, groupes scolaires régionaux", "Plateforme d’exploitation scolaire, plus seulement un ENT"],
        ],
        [3.2, 7.0, 7.0],
    )

    # ---- 12 ----
    h1("12  ·  Partenariats recherchés et prochaines étapes")
    h2("Ce que nous proposons, concrètement")
    para("Selon le profil du destinataire, plusieurs portes d’entrée existent. Elles ne s’excluent pas.")
    add_table(
        ["Vous êtes…", "Ce que nous vous proposons", "Ce que vous y gagnez"],
        [
            ["Promoteur ou réseau d’écoles", "Statut de pilote fondateur, tarif préférentiel, co-construction du déploiement", "Un outil en avance, une voix sur la feuille de route, un avantage vis-à-vis des familles"],
            ["Association de promoteurs", "Partenariat de déploiement, sessions de démonstration, offre réseau", "Un service concret pour vos membres, une modernisation visible"],
            ["Opérateur telecom / Mobile Money", "Intégration des paiements de scolarité dans Sulungukutu", "Un flux récurrent, ancré dans un usage quotidien des familles"],
            ["Institution financière / fintech éducative", "Données de recouvrement (avec accord) et écoles mieux lisibles", "Moins de risque sur les crédits rentrée, un canal d’écoles structurées"],
            ["Investisseur / business angel", "Entrée au capital ou financement d’amorçage go-to-market", "Un produit déjà construit, un marché concentré, un modèle SaaS + paiements"],
            ["Collaborateur clé", "Aventure de déploiement, commercial, succès client, partnerships", "Une entreprise naissante sur un besoin réel, pas une idée de salon"],
        ],
        [4.2, 6.6, 6.4],
    )
    h2("Emploi indicatif d’un financement d’amorçage")
    para("À titre de cadrage — le montant exact et les modalités se discutent — un enveloppe de l’ordre de 25 à 40 millions FCFA sur 12 à 18 mois permettrait d’industrialiser sans diluer l’attention :", italic=False)
    add_table(
        ["Poste", "Part indicative", "À quoi ça sert"],
        [
            ["Commercialisation et onboarding", "≈ 40 %", "Demos, déplacements Brazza / PNR, formation, success des pilotes"],
            ["Paiements réels et SMS", "≈ 25 %", "Intégrations Airtel / MTN, passerelle SMS, tests de recouvrement"],
            ["Produit (app parents, durcissement)", "≈ 20 %", "Aller chercher les familles là où elles sont : le téléphone"],
            ["Infra, juridique, support", "≈ 15 %", "Hébergement, contrats, assistance rentrée"],
        ],
        [5.2, 3.6, 8.4],
    )
    h2("Les trois prochaines réunions qui comptent")
    bullets([
        "Une démonstration live de 45 minutes, sur l’école de démonstration, dans les cinq rôles.",
        "La visite d’un établissement candidat pilote : import des élèves, paramétrage, formation d’une matinée.",
        "Un protocole d’accord (pilote, distribution, ou discussion d’investissement) avec un calendrier de rentrée.",
    ])
    callout(
        "Invitation",
        "Si vous dirigez une école, un réseau, un opérateur ou un capital, Sulungukutu n’a pas besoin que vous « y croyiez » en l’air. Asseyez-vous devant le produit. Demandez le tableau des impayés, un bulletin, l’espace d’un parent. Puis décidez. C’est la seule séquence honnête — et c’est celle que nous proposons.",
        fill="EEF2FF",
        accent=INDIGO_HEX,
    )

    h1("Contact")
    para("Kassy Gloire Exaucé")
    para("Fondateur — Sulungukutu", bold=True, after=4)
    para("Plateforme de gestion scolaire pour l’Afrique centrale")
    para("Document de présentation partenaires  ·  août 2026  ·  usage confidentiel")
    para("Ce document ne constitue pas une offre au public. Les tarifs, projections et scénarios de recouvrement sont fournis à titre illustratif pour éclairer une discussion. Toute collaboration fera l’objet d’un accord écrit.", italic=True, size=9, color=MUTED)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_PATH))


class DossierPDF(FPDF):
    def __init__(self):
        super().__init__(format="A4", unit="mm")
        self.set_margins(18, 22, 18)
        self.set_auto_page_break(auto=True, margin=18)
        self.alias_nb_pages()
        if FONT_REG.exists():
            self.add_font("Body", "", str(FONT_REG))
            self.add_font("Body", "B", str(FONT_BOLD if FONT_BOLD.exists() else FONT_REG))
            self.add_font("Body", "I", str(FONT_ITALIC if FONT_ITALIC.exists() else FONT_REG))
            self.add_font("Body", "BI", str(FONT_BI if FONT_BI.exists() else FONT_REG))
            title_reg = FONT_TITLE_REG if FONT_TITLE_REG.exists() else FONT_REG
            title_b = FONT_TITLE if FONT_TITLE.exists() else FONT_BOLD
            self.add_font("Title", "", str(title_reg))
            self.add_font("Title", "B", str(title_b))
            self.font_body = "Body"
            self.font_title = "Title"
        else:
            self.font_body = "Helvetica"
            self.font_title = "Helvetica"

    def header(self):
        if self.page_no() == 1:
            return
        self.set_y(10)
        self.set_font(self.font_title, "B", 9)
        self.set_text_color(79, 70, 229)
        self.cell(40, 6, "SULUNGUKUTU")
        self.set_font(self.font_body, "", 9)
        self.set_text_color(87, 83, 78)
        self.cell(0, 6, "   Dossier de présentation partenaires", align="L")
        self.set_draw_color(79, 70, 229)
        self.set_line_width(0.4)
        self.line(18, 17, 192, 17)
        self.set_y(22)

    def footer(self):
        if self.page_no() == 1:
            return
        self.set_y(-14)
        self.set_draw_color(231, 229, 228)
        self.set_line_width(0.2)
        self.line(18, self.get_y(), 192, self.get_y())
        self.set_y(-12)
        self.set_font(self.font_body, "", 8)
        self.set_text_color(87, 83, 78)
        self.cell(120, 6, "Confidentiel  ·  Ne pas diffuser sans autorisation")
        self.cell(0, 6, f"Page {self.page_no()} / {{nb}}", align="R")

    def h1(self, text: str):
        self.ln(4)
        if self.get_y() > 250:
            self.add_page()
        self.set_font(self.font_title, "B", 15)
        self.set_text_color(79, 70, 229)
        self.multi_cell(0, 8, text)
        y = self.get_y()
        self.set_draw_color(224, 231, 255)
        self.set_line_width(0.7)
        self.line(18, y + 0.5, 192, y + 0.5)
        self.ln(4)

    def h2(self, text: str):
        if self.get_y() > 255:
            self.add_page()
        self.ln(2)
        self.set_font(self.font_title, "B", 12)
        self.set_text_color(34, 31, 29)
        self.multi_cell(0, 7, text)
        self.ln(1)

    def body(self, text: str, italic=False):
        if self.get_y() > 262:
            self.add_page()
        style = "I" if italic else ""
        self.set_font(self.font_body, style, 10.5)
        self.set_text_color(28, 25, 23)
        self.multi_cell(0, 5.6, text)
        self.ln(1.5)

    def muted(self, text: str):
        self.set_font(self.font_body, "I", 9.5)
        self.set_text_color(87, 83, 78)
        self.multi_cell(0, 5.2, text)
        self.ln(1.5)

    def bullets(self, items: list[str]):
        for item in items:
            if self.get_y() > 265:
                self.add_page()
            x = self.get_x()
            y = self.get_y()
            self.set_fill_color(79, 70, 229)
            self.ellipse(x + 0.4, y + 1.8, 2.0, 2.0, "F")
            self.set_xy(x + 5, y)
            self.set_font(self.font_body, "", 10.5)
            self.set_text_color(28, 25, 23)
            self.multi_cell(169, 5.5, item)
            self.ln(0.8)

    def callout(self, title: str, body: str, fill=(238, 242, 255), border=(79, 70, 229)):
        self.ln(1)
        self.set_font(self.font_body, "B", 9.5)
        title_h = 6
        self.set_font(self.font_body, "", 10)
        body_h = self._lines_height(body, 164, 5.3)
        h = 8 + title_h + body_h
        if self.get_y() + h > 270:
            self.add_page()
        x, y = 18, self.get_y()
        self.set_fill_color(*fill)
        self.set_draw_color(*border)
        self.set_line_width(0.5)
        self.rect(x, y, 174, h, "FD")
        self.set_xy(x + 5, y + 3)
        self.set_font(self.font_body, "B", 9.5)
        self.set_text_color(*border)
        self.cell(164, 5, title)
        self.set_xy(x + 5, y + 9)
        self.set_font(self.font_body, "", 10)
        self.set_text_color(28, 25, 23)
        self.multi_cell(164, 5.3, body)
        self.set_y(y + h + 3)

    def _lines_height(self, text: str, width: float, line_h: float) -> float:
        self.set_font(self.font_body, "", 10)
        lines = self.multi_cell(width, line_h, text, dry_run=True, output="LINES")
        return max(1, len(lines)) * line_h

    def kpi(self, items: list[tuple[str, str, str]]):
        n = len(items)
        w = 174 / n
        h = 28
        if self.get_y() + h > 270:
            self.add_page()
        x0, y0 = 18, self.get_y()
        for i, (k, v, sub) in enumerate(items):
            x = x0 + i * w
            fill = (238, 242, 255) if i % 2 == 0 else (255, 251, 235)
            self.set_fill_color(*fill)
            self.set_draw_color(224, 231, 255)
            self.rect(x + 0.8, y0, w - 1.6, h, "FD")
            self.set_xy(x + 3, y0 + 3)
            self.set_font(self.font_body, "B", 7.5)
            self.set_text_color(79, 70, 229)
            self.cell(w - 6, 4, k.upper())
            self.set_xy(x + 3, y0 + 9)
            self.set_font(self.font_title, "B", 12)
            self.set_text_color(34, 31, 29)
            self.cell(w - 6, 6, v)
            self.set_xy(x + 3, y0 + 17)
            self.set_font(self.font_body, "", 7.5)
            self.set_text_color(87, 83, 78)
            self.multi_cell(w - 6, 3.6, sub)
        self.set_y(y0 + h + 4)

    def table(self, headers: list[str], rows: list[list[str]], widths: list[float]):
        line_h = 4.8
        header_h = 8

        def row_h(vals: list[str], font_size=9) -> float:
            self.set_font(self.font_body, "", font_size)
            mh = header_h if vals is headers else line_h
            for val, w in zip(vals, widths):
                lines = self.multi_cell(w - 3.2, line_h, val, dry_run=True, output="LINES")
                mh = max(mh, len(lines) * line_h + 2.4)
            return mh

        if self.get_y() + header_h + 10 > 270:
            self.add_page()
        x0 = 18
        y = self.get_y()
        # header
        hh = row_h(headers, 8)
        self.set_fill_color(34, 31, 29)
        self.set_draw_color(34, 31, 29)
        self.rect(x0, y, 174, hh, "F")
        x = x0
        for h, w in zip(headers, widths):
            self.set_xy(x + 1.6, y + 1.4)
            self.set_font(self.font_body, "B", 8)
            self.set_text_color(255, 255, 255)
            self.multi_cell(w - 3.2, line_h, h)
            x += w
        y += hh
        for r_i, row in enumerate(rows):
            rh = row_h(row, 9)
            if y + rh > 272:
                self.add_page()
                y = self.get_y()
                self.set_fill_color(34, 31, 29)
                self.rect(x0, y, 174, hh, "F")
                x = x0
                for h, w in zip(headers, widths):
                    self.set_xy(x + 1.6, y + 1.4)
                    self.set_font(self.font_body, "B", 8)
                    self.set_text_color(255, 255, 255)
                    self.multi_cell(w - 3.2, line_h, h)
                    x += w
                y += hh
            fill = (238, 242, 255) if r_i % 2 == 0 else (255, 255, 255)
            self.set_fill_color(*fill)
            self.set_draw_color(231, 229, 228)
            self.set_line_width(0.2)
            self.rect(x0, y, 174, rh, "FD")
            x = x0
            for val, w in zip(row, widths):
                self.set_xy(x + 1.6, y + 1.2)
                self.set_font(self.font_body, "", 9)
                self.set_text_color(28, 25, 23)
                self.multi_cell(w - 3.2, line_h, val)
                x += w
            y += rh
        self.set_y(y + 4)


def build_pdf() -> None:
    pdf = DossierPDF()
    pdf.set_title("Sulungukutu — Dossier de présentation partenaires")
    pdf.set_author("Kassy Gloire Exaucé")
    pdf.set_creator("Sulungukutu")
    pdf.set_subject("Plateforme de gestion scolaire — dossier partenaires et investisseurs")

    # Cover
    pdf.add_page()
    pdf.set_auto_page_break(auto=False)
    pdf.set_fill_color(34, 31, 29)
    pdf.rect(0, 0, 210, 297, "F")
    pdf.set_fill_color(79, 70, 229)
    pdf.rect(0, 0, 7, 297, "F")
    pdf.set_fill_color(245, 158, 11)
    pdf.rect(7, 0, 1.2, 297, "F")

    if LOGO.exists():
        pdf.image(str(LOGO), x=22, y=22, w=22)

    pdf.set_xy(22, 52)
    pdf.set_font(pdf.font_body, "B", 10)
    pdf.set_text_color(245, 158, 11)
    pdf.cell(0, 6, "DOCUMENT CONFIDENTIEL  ·  PARTENAIRES ET INVESTISSEURS")

    pdf.set_xy(22, 68)
    pdf.set_font(pdf.font_title, "B", 36)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 14, "SULUNGUKUTU")

    pdf.set_xy(22, 86)
    pdf.set_font(pdf.font_title, "", 18)
    pdf.set_text_color(199, 210, 254)
    pdf.cell(0, 10, "Dossier de présentation")

    pdf.set_draw_color(245, 158, 11)
    pdf.set_line_width(1.1)
    pdf.line(22, 104, 78, 104)

    pdf.set_xy(22, 112)
    pdf.set_font(pdf.font_body, "", 13)
    pdf.set_text_color(255, 255, 255)
    pdf.multi_cell(166, 7, "La plateforme de gestion scolaire conçue pour les établissements d’Afrique centrale.")

    pdf.set_xy(22, 132)
    pdf.set_font(pdf.font_body, "", 11)
    pdf.set_text_color(214, 211, 209)
    pdf.multi_cell(166, 6.2, "Un seul outil pour piloter la pédagogie, la scolarité et la relation avec les familles — et rendre enfin visible l’argent, les absences et les résultats.")

    meta = [
        ("Promoteur", "Kassy Gloire Exaucé"),
        ("Produit", "SaaS multi-établissements — version 1.0 opérationnelle"),
        ("Marché prioritaire", "Écoles privées du Congo-Brazzaville, puis CEMAC"),
        ("Document", "Dossier partenaires  ·  août 2026"),
    ]
    y = 168
    for k, v in meta:
        pdf.set_xy(22, y)
        pdf.set_font(pdf.font_body, "B", 10)
        pdf.set_text_color(245, 158, 11)
        pdf.cell(42, 6, k)
        pdf.set_font(pdf.font_body, "", 10)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(0, 6, v)
        y += 9

    pdf.set_xy(22, 250)
    pdf.set_font(pdf.font_title, "B", 12)
    pdf.set_text_color(199, 210, 254)
    pdf.multi_cell(166, 7, "« Le bulletin arrive avant la fin du trimestre. »")
    pdf.set_xy(22, 272)
    pdf.set_font(pdf.font_body, "", 9)
    pdf.set_text_color(168, 162, 158)
    pdf.cell(0, 5, "Usage interne et partenaires  ·  Ne pas diffuser sans autorisation")

    # Body
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()

    pdf.h1("Sommaire")
    pdf.body("Ce document présente Sulungukutu dans son ensemble : le problème, le produit, le modèle économique, l’impact opérationnel et les conditions d’un partenariat ou d’un investissement.")
    for num, title in SOMMAIRE:
        pdf.set_font(pdf.font_body, "B", 11)
        pdf.set_text_color(79, 70, 229)
        pdf.cell(14, 7, num)
        pdf.set_font(pdf.font_body, "", 11)
        pdf.set_text_color(28, 25, 23)
        pdf.cell(0, 7, title, new_x="LMARGIN", new_y="NEXT")

    pdf.h1("01  ·  L’essentiel")
    pdf.body("Sulungukutu est une plateforme numérique de gestion scolaire, pensée pour les établissements d’Afrique subsaharienne et d’abord calibrée sur le système éducatif congolais. Elle replace le trio encore dominant — cahiers, fichiers Excel, groupes WhatsApp — par une source unique de vérité, partagée entre la direction, les enseignants, les parents et les élèves.")
    pdf.body("Le produit n’est pas une intention. La version 1.0 est développée, déployable, et couvre déjà le cycle de vie d’une année scolaire : inscriptions, classes, notes, présences, bulletins trimestriels, mensualités, messagerie, annonces, exports et journal d’audit. L’architecture est multi-établissements : un promoteur peut piloter plusieurs écoles depuis un même compte.")
    pdf.kpi([
        ("Produit", "Version 1.0", "Opérationnelle, prête pour pilotes"),
        ("Utilisateurs", "5 espaces", "Direction, enseignants, parents, élèves, super-admin"),
        ("Cœur métier", "Pédagogie + scolarité", "Notes, bulletins, présences, paiements"),
        ("Ambition", "SaaS régional", "Congo d’abord, CEMAC ensuite"),
    ])
    pdf.callout(
        "La promesse, en une phrase",
        "Donner à chaque établissement une vision en temps réel de ses élèves, de ses résultats et de sa trésorerie — et redonner aux familles le droit de savoir, sans attendre le conseil de classe.",
    )
    pdf.body("Nous présentons ce dossier à des collaborateurs, des partenaires stratégiques et des investisseurs potentiels. L’objectif n’est pas de vendre un rêve : c’est de montrer un outil déjà construit, un marché réel, un modèle commercial simple, et une fenêtre d’entrée encore ouverte.")

    pdf.h1("02  ·  Le problème que nous résolvons")
    pdf.h2("Une école privée congolaise gagne ou perd sur la gestion, pas seulement sur la pédagogie")
    pdf.body("Au Congo-Brazzaville, le privé n’est plus un complément. Selon les données publiques récentes, environ 41 % des élèves de l’enseignement général fréquentent un établissement privé (PASEC / MEPPSA). En primaire, près d’un établissement sur deux est privé, et près de la moitié des élèves y sont scolarisés. En 2026, 1 192 établissements privés d’enseignement général étaient soumis à la commission d’agrément.")
    pdf.body("Ces écoles vivent pourtant encore, pour la plupart, dans un régime artisanal :")
    pdf.bullets([
        "Les notes circulent dans des cahiers, puis sont recopiées, puis recalculées à la main au moment des bulletins.",
        "Les présences se marquent le matin et s’évaporent le soir : le parent apprend une absence trop tard, parfois jamais.",
        "La scolarité se suit sur un registre ou un Excel. Les acomptes, les exonérations et les relances se perdent. L’impayé n’est pas un incident : c’est un trou de trésorerie structurel.",
        "La communication passe par WhatsApp, sans traçabilité, sans ciblage, sans preuve.",
        "Le promoteur pilote à l’intuition. Il découvre les décrochages, les classes en difficulté et les retards de paiement quand il est trop tard pour agir.",
    ])
    pdf.body("Ce n’est pas un détail opérationnel. C’est le cœur du risque économique de l’école privée. Les institutions financières locales proposent précisément des crédits aux promoteurs « pour payer les salaires sans attendre que tous les parents aient réglé ». Autrement dit : le recouvrement est le talon d’Achille du secteur — et personne n’a encore industrialisé sa visibilité.")
    pdf.table(
        ["Acteur", "Ce qu’il vit aujourd’hui", "Ce qu’il devrait vivre"],
        [
            ["Promoteur / Directeur", "Pilotage à l’aveugle, trésorerie stressée, image artisanale", "Tableau de bord du jour : effectifs, absences, impayés, moyennes"],
            ["Secrétariat", "Saisies multiples, relances papier, bulletins de dernière minute", "Une base unique, des exports, des bulletins générés"],
            ["Enseignant", "Cahier de notes, appels répétitifs, peu de feedback", "Saisie simple, présences, classes sous les yeux"],
            ["Parent", "Découvre les résultats trop tard, doute sur les paiements", "Suivi de l’enfant, bulletin publié, historique de scolarité"],
            ["Élève", "Peu de visibilité sur son propre parcours", "Notes, bulletin, présences, emploi du temps"],
        ],
        [38, 68, 68],
    )
    pdf.callout(
        "Le coût invisible",
        "Chaque heure perdue au secrétariat, chaque mensualité « oubliée », chaque bulletin publié en retard, chaque parent qui retire son enfant faute de confiance : ce n’est pas de l’intendance. C’est du chiffre d’affaires, de la réputation et de la qualité pédagogique qui s’évaporent.",
        fill=(255, 251, 235),
        border=(217, 119, 6),
    )

    pdf.h1("03  ·  Un marché prêt, encore peu équipé")
    pdf.h2("Un socle national concentré et adressable")
    pdf.body("Le marché congolais a une caractéristique rare pour un investisseur : il est à la fois suffisamment grand pour construire une entreprise, et suffisamment concentré pour être attaqué avec une équipe courte.")
    pdf.table(
        ["Signal de marché", "Donnée", "Implication pour Sulungukutu"],
        [
            ["Écoles privées d’enseignement général", "1 192 dossiers d’agrément (2026)", "Cible B2B identifiable, organisée en associations"],
            ["Poids du privé", "≈ 41 % des élèves (général) ; ≈ 49 % en primaire", "Le privé est central, pas périphérique"],
            ["Géographie", "89 % des écoles privées à Brazzaville et Pointe-Noire", "Deux villes suffisent pour un premier réseau dense"],
            ["Ratio d’encadrement privé", "≈ 27 élèves / enseignant (vs 48 dans le public)", "Des écoles déjà plus « gérables » numériquement"],
            ["Pression réglementaire", "Agrément, mise en conformité, modernisation", "La traçabilité devient un argument, pas un luxe"],
            ["Digitalisation en cours", "Outils locaux encore émergents, sessions promoteurs", "La question n’est plus « faut-il digitaliser ? » mais « avec qui ? »"],
        ],
        [46, 58, 70],
    )
    pdf.h2("Une fenêtre, pas un océan rouge")
    pdf.body("Des solutions existent — registres papier, Excel, quelques logiciels locaux, et des suites étrangères inadaptées au Mobile Money, aux mensualités de septembre à mai, et au référentiel national. Le marché n’est pas vide. Il est fragmenté, peu saturé, et en demande de recouvrement autant que de pédagogie.")
    pdf.body("La CEMAC offre ensuite une expansion naturelle : mêmes langues de travail, même culture des trimestres et des coefficients, même pénétration du Mobile Money, mêmes promoteurs multi-sites. Le Congo n’est pas le plafond. C’est la tête de pont.")
    pdf.callout(
        "Lecture investisseur",
        "Nous ne promettons pas un « marché de plusieurs milliards » pour impressionner. Nous montrons un premier pays où plus de mille écoles privées existent, où la moitié des élèves du primaire sont déjà dans le privé, où deux villes concentrent la demande, et où le paiement des familles est le nerf de la guerre. C’est un marché vendable, pas un slogan.",
    )

    pdf.h1("04  ·  La solution Sulungukutu")
    pdf.body("Sulungukutu est un logiciel en ligne (SaaS), accessible depuis un navigateur, sans installation lourde dans l’école. Chaque établissement dispose de son espace isolé, de son identité visuelle (logo, couleur), de son année scolaire, et de ses utilisateurs. Un même adulte — promoteur, enseignant, parent — peut appartenir à plusieurs écoles et basculer d’un établissement à l’autre.")
    pdf.h2("Ce qui la rend différente d’un « logiciel scolaire » générique")
    pdf.bullets([
        "Elle est née du terrain congolais : trimestres T1 / T2 / T3, mentions, coefficients, séries du lycée (A, C, D, TI), mensualités de septembre à mai, paiements espèces, Mobile Money, virement, chèque, exonérations et acomptes.",
        "Le référentiel pédagogique national (CP1 → Terminale) est intégré. L’école n’a pas à recréer le programme : elle active les niveaux et les matières dont elle a besoin.",
        "Elle relie pédagogie et argent. Un bulletin, une absence, une mensualité impayée ne vivent plus dans trois univers séparés.",
        "Elle parle aux parents, pas seulement au secrétariat. Le parent voit les résultats publiés, les présences, l’historique de scolarité, et peut échanger avec l’école.",
        "Elle est conçue pour des utilisateurs sans e-mail personnel : connexion par identifiant, e-mail ou téléphone ; réinitialisation du mot de passe en présentiel par l’administration.",
    ])
    pdf.h2("Cinq espaces, une seule base")
    pdf.table(
        ["Espace", "Pour qui", "Ce qu’il permet"],
        [
            ["Administration", "Directeur, promoteur, secrétariat", "Pilotage, élèves, enseignants, notes, bulletins, paiements, emplois du temps, annonces, exports, audit"],
            ["Enseignant", "Professeurs", "Notes (devoir, contrôle, examen, interro), présences, bulletins, emploi du temps, messagerie"],
            ["Parent", "Pères, mères, tuteurs", "Suivi des enfants, bulletins publiés, scolarité, présences, messages"],
            ["Élève", "Collégiens et lycéens", "Notes, bulletin, présences, emploi du temps, messagerie"],
            ["Super-admin", "Opérateur / groupe scolaire", "Toutes les écoles, accès, statistiques consolidées"],
        ],
        [34, 48, 92],
    )

    pdf.h1("05  ·  L’application dans son ensemble")
    pdf.body("Sulungukutu n’est pas un module isolé. C’est une suite complète, déjà assemblée, qui couvre le quotidien d’un établissement du premier jour de rentrée jusqu’à la publication des bulletins.")
    pdf.h2("Pilotage")
    pdf.body("Le tableau de bord direction affiche en temps réel les effectifs, le taux de présence du jour, les impayés du mois, les moyennes par classe, la répartition des mentions et l’activité récente. Des alertes intelligentes et un onboarding guidé accompagnent les nouvelles écoles. Le promoteur cesse de « demander au secrétariat » : il voit.")
    pdf.h2("Organisation pédagogique")
    pdf.bullets([
        "Niveaux, classes, matières, coefficients, volumes horaires.",
        "Activation du référentiel national (primaire, collège, lycée) sans ressaisie.",
        "Emplois du temps, avec détection des conflits d’un enseignant avant création d’un créneau.",
        "Calendrier de l’établissement et transfert d’élève d’une classe à une autre.",
    ])
    pdf.h2("Élèves et familles")
    pdf.bullets([
        "Inscription d’un élève avec création automatique du compte parent.",
        "Lien père / mère / tuteur, y compris un parent déjà connu pour un autre enfant.",
        "Import CSV pour les rentrées, annuaire des identifiants, invitations enseignants.",
        "Identité visuelle de l’école : nom, logo, couleur d’accent.",
    ])
    pdf.h2("Notes, présences, bulletins")
    pdf.bullets([
        "Saisie unitaire ou en masse : devoir, contrôle, examen, interrogation.",
        "Appel par matière et par jour : présent, absent, retard — avec statistiques.",
        "Génération des bulletins trimestriels, moyennes et mentions, publication contrôlée, archivage, PDF.",
        "Le parent et l’élève ne voient le bulletin que lorsqu’il est publié : la direction garde la main.",
    ])
    pdf.h2("Scolarité et trésorerie")
    pdf.bullets([
        "Neuf mensualités par année scolaire : payé, partiel, impayé, exonéré, annulé.",
        "Journal de transactions : espèces au guichet, Mobile Money, virement, chèque.",
        "Acomptes, annulation, historique, liste des impayés du mois, relances e-mail.",
        "Export Excel du recouvrement — le document que le promoteur emporte à la banque.",
    ])
    pdf.h2("Communication, preuve et conformité")
    pdf.bullets([
        "Messagerie interne, annonces ciblées, notifications en temps réel.",
        "Journal d’audit : qui a modifié un paiement, publié un bulletin, invité un utilisateur.",
        "Exports Excel, isolation stricte entre écoles, rôles et sessions révocables.",
    ])

    pdf.h1("06  ·  Ce que Sulungukutu réduit, ce qu’elle améliore")
    pdf.h2("Elle réduit")
    pdf.table(
        ["Elle réduit…", "Comment", "Effet pour l’école"],
        [
            ["Le temps administratif", "Bulletins générés, listes exportées, relances, import de rentrée", "Le secrétariat cesse d’être un goulot d’étranglement"],
            ["Les impayés « invisibles »", "Statut de chaque mensualité, retardataires, historique des acomptes", "La trésorerie redevient lisible, donc actionnable"],
            ["Les erreurs de moyennes", "Calcul automatisé, mentions, publication contrôlée", "Moins de contestations, plus de crédibilité"],
            ["Le papier et les files", "Espaces parent / élève, PDF, messagerie, annonces", "Moins d’impressions, moins d’allers-retours au secrétariat"],
            ["La perte d’information", "Une base unique à la place de cahiers et d’Excel dispersés", "Plus de « le fichier est chez Madame X »"],
            ["Le délai parent–école", "Présences, notes publiées, notifications", "La confiance se construit en continu, pas au trimestre"],
            ["Le risque de contestation", "Journal d’audit, transactions horodatées, rôles", "Chaque acte sensible laisse une trace"],
            ["Le travail en double", "Un élève, un parent, une note, un paiement : saisis une fois", "Moins de fatigue, moins d’incohérences"],
        ],
        [42, 68, 64],
    )
    pdf.h2("Elle améliore")
    pdf.table(
        ["Elle améliore…", "Pour qui", "Pourquoi c’est décisif"],
        [
            ["Le recouvrement de la scolarité", "Promoteur", "C’est le salaire des enseignants et la survie de l’école"],
            ["La transparence", "Parents", "Un parent informé reste, recommande, paie mieux"],
            ["Le suivi pédagogique", "Enseignants & direction", "On voit les classes qui décrochent avant l’examen"],
            ["La qualité des décisions", "Direction", "KPI du jour plutôt que réunion de crise"],
            ["L’image de l’établissement", "Tout le monde", "Une école digitale inspire confiance aux familles urbaines"],
            ["La fidélisation des élèves", "Promoteur", "Partir coûte plus cher qu’un abonnement logiciel"],
            ["Les conditions de travail", "Personnel", "Moins de paperasse, plus de métier"],
            ["La conformité et l’agrément", "Promoteur", "Traçabilité des effectifs, des actes et des paiements"],
        ],
        [50, 38, 86],
    )
    pdf.h2("Scénario illustratif — collège privé de 320 élèves à Brazzaville")
    pdf.muted("Les chiffres ci-dessous sont un scénario de travail, pas une garantie de résultat. Ils servent à rendre concret le retour sur investissement pour un promoteur.")
    pdf.table(
        ["Hypothèse", "Ordre de grandeur"],
        [
            ["Effectif", "320 élèves"],
            ["Mensualité moyenne", "22 000 FCFA"],
            ["Recette théorique mensuelle", "7,04 millions FCFA"],
            ["Retards / impayés chroniques (hypothèse 18 %)", "≈ 1,27 million FCFA « gelés » chaque mois"],
            ["Si Sulungukutu ne fait gagner que 6 points de recouvrement", "≈ 422 000 FCFA / mois  ·  ≈ 3,8 millions FCFA / an"],
            ["Coût logiciel (offre Professionnel, indicatif)", "45 000 FCFA / mois  ·  540 000 FCFA / an"],
            ["Ordre de grandeur du multiple", "Le seul recouvrement peut couvrir le logiciel plusieurs fois — avant le temps gagné au secrétariat"],
        ],
        [92, 82],
    )
    pdf.callout(
        "Ce que l’investisseur doit retenir",
        "Sulungukutu ne se vend pas comme un « plus pédagogique ». Elle se vend comme un instrument de trésorerie, de réputation et de temps. C’est pour cela qu’un promoteur peut justifier la dépense, et qu’un partenaire peut y voir un revenu récurrent défendable.",
        fill=(236, 253, 245),
        border=(5, 150, 105),
    )

    pdf.h1("07  ·  Modèle commercial et go-to-market")
    pdf.h2("Comment Sulungukutu gagne de l’argent")
    pdf.body("Le modèle est un abonnement SaaS facturé à l’établissement, selon la taille. Simple à expliquer, récurrent, prévisible. L’école n’achète pas une licence figée : elle paie un service vivant (mises à jour, support, hébergement, nouvelles fonctions).")
    pdf.table(
        ["Offre", "Cible", "Tarif indicatif", "Inclus"],
        [
            ["Essentiel", "Jusqu’à 200 élèves", "20 000 FCFA / mois", "Modules cœur, 1 établissement, support e-mail"],
            ["Professionnel", "200 à 600 élèves", "45 000 FCFA / mois", "Essentiel + imports, exports, onboarding assisté"],
            ["Groupe scolaire", "Multi-sites / + 600 élèves", "Sur devis + 80 FCFA / élève / mois", "Super-admin, consolidé, accompagnement dédié"],
            ["Mise en service", "Toutes les offres", "50 000 FCFA (une fois)", "Paramétrage, import rentrée, formation direction + secrétariat"],
        ],
        [34, 42, 42, 56],
    )
    pdf.body("Remise de 15 % en paiement annuel anticipé — aligné sur le rythme de trésorerie de la rentrée. Les tarifs sont une grille de travail, destinée à ancrer la discussion.")
    pdf.h2("Revenus additionnels, dans un second temps")
    pdf.bullets([
        "Commission discrète sur les paiements Mobile Money en production.",
        "SMS de relance et d’absence, facturés au volume ou inclus dans une offre supérieure.",
        "Modules avancés : paie du personnel, cantine, transport, bibliothèque.",
        "Offre « réseau » pour associations de promoteurs ou groupes confessionnels.",
    ])
    pdf.h2("Comment nous commercialisons")
    pdf.body("Le cycle de vente scolaire est saisonnier. On signe entre mai et septembre. On enchante à la rentrée. On renouvelle l’année suivante si les bulletins du T1 sont sortis sans douleur et si le recouvrement de novembre est lisible.")
    pdf.table(
        ["Phase", "Horizon", "Objectif", "Moyens"],
        [
            ["Pilotes", "6 à 9 mois", "10 à 15 écoles à Brazzaville et Pointe-Noire", "Démo live, tarif pilote, accompagnement, preuves chiffrées"],
            ["Réseau", "12 à 24 mois", "50+ établissements au Congo", "Associations de promoteurs, multi-sites, bouche-à-oreille"],
            ["Distribution", "En parallèle", "Canal telco et financier", "Airtel / MTN, microfinance scolaire, agréments"],
            ["Expansion CEMAC", "24 à 36 mois", "Gabon, RDC, Cameroun", "Même produit, référentiels additionnels, relais locaux"],
        ],
        [32, 30, 52, 60],
    )
    pdf.h2("Ordre de grandeur de revenus (scénario prudent)")
    pdf.muted("Projections indicatives, hors commission Mobile Money. Elles illustrent le levier d’un modèle récurrent, même avec un rythme commercial réaliste.")
    pdf.table(
        ["Horizon", "Écoles actives", "Panier moyen mensuel", "ARR indicatif"],
        [
            ["Année 1", "25", "35 000 FCFA", "≈ 10,5 millions FCFA"],
            ["Année 2", "70", "38 000 FCFA", "≈ 31,9 millions FCFA"],
            ["Année 3", "140", "42 000 FCFA", "≈ 70,6 millions FCFA"],
        ],
        [36, 42, 48, 48],
    )
    pdf.body("À 140 écoles, Sulungukutu n’a encore adressé qu’environ 12 % du seul vivier privé congolais d’enseignement général. Le plafond local n’est pas atteint. L’expansion régionale n’est même pas nécessaire pour justifier une première brique d’entreprise — elle en accélère simplement l’ampleur.")

    pdf.h1("08  ·  Pourquoi un établissement paie — et pourquoi un investisseur entre")
    pdf.h2("Le motif d’achat de l’école")
    pdf.body("Un promoteur n’achète pas « du digital ». Il achète trois choses qu’il peut défendre devant son conseil, son associée, ou son banquier :")
    pdf.bullets([
        "Voir l’argent : qui a payé, qui n’a pas payé, depuis quand, pour quel enfant.",
        "Tenir la promesse aux parents : un bulletin à l’heure, une absence signalée, une école qui a l’air tenue.",
        "Diminuer la charge du secrétariat sans embaucher une personne de plus.",
    ])
    pdf.h2("Le motif d’entrée de l’investisseur ou du partenaire")
    pdf.bullets([
        "Le risque produit est déjà largement derrière nous : la plateforme existe, les cinq rôles existent, le métier école est modélisé.",
        "Le revenu est récurrent, collant (on ne change pas de logiciel en janvier), et saisonnier de façon prévisible.",
        "Le paiement de la scolarité est un flux réel, mensuel, massif. S’y brancher crée un actif difficile à déloger.",
        "Le marché est localement concentré : on peut gagner une ville, puis un pays, sans brûler du capital en acquisition mondiale.",
        "L’impact est lisible : écoles mieux gérées, parents mieux informés, enseignants moins submergés.",
        "La barrière se construit avec le temps : données de l’année, habitudes des enseignants, comptes parents, historique de paiements.",
    ])
    pdf.callout(
        "Ce n’est pas un projet « à construire si vous financez »",
        "Trop de dossiers africains demandent de l’argent pour commencer à coder. Ici, l’argent — et les alliances — servent à vendre, à accompagner, à connecter le Mobile Money réel, et à transformer un produit abouti en standard de fait des écoles privées congolaises.",
    )

    pdf.h1("09  ·  Avantage concurrentiel")
    pdf.table(
        ["Alternative", "Limite", "Ce que Sulungukutu oppose"],
        [
            ["Cahiers + Excel + WhatsApp", "Fragile, non partagé, non auditable, lent", "Une source de vérité, des rôles, un historique"],
            ["Logiciel étranger (Pronote, etc.)", "Cher, mal calé sur le programme local, peu Mobile Money", "Référentiel congolais, mensualités, prix africain"],
            ["Outil local partiel", "Souvent un module (caisse ou notes), peu d’espace parent moderne", "Suite complète, cinq espaces, multi-écoles"],
            ["Développement interne de l’école", "Coûteux, jamais fini, dépendant d’une personne", "Produit maintenu, sécurisé, déjà prêt"],
        ],
        [50, 62, 62],
    )
    pdf.h2("Avantages que nous renforçons volontairement")
    pdf.bullets([
        "Référentiel national intégré — plus on l’enrichit, plus un concurrent générique est hors-jeu.",
        "Multi-établissements natif — les promoteurs sérieux ont plusieurs sites.",
        "Couche parents — c’est elle qui rend le logiciel collant, pas seulement le secrétariat.",
        "Journal d’audit et isolation des données — arguments de confiance pour les groupes et, demain, pour le régulateur.",
        "Distribution par les réseaux de promoteurs — le commercial de l’éducation se gagne en salle, pas en bannière.",
    ])

    pdf.h1("10  ·  Socle technique et sécurité")
    pdf.body("Un partenaire technologique ou un investisseur due-diligence doit pouvoir vérifier que Sulungukutu n’est pas un prototype fragile. L’architecture est celle d’un produit moderne, déjà pensé pour plusieurs écoles et pour la production.")
    pdf.table(
        ["Couche", "Choix", "Pourquoi c’est un atout"],
        [
            ["Application", "Next.js 14, cinq espaces, interface soignée", "Expérience contemporaine, mobile-friendly"],
            ["API", "GraphQL, temps réel (notifications, messages, paiements)", "Un écran direction qui se met à jour sans recharger"],
            ["Données", "PostgreSQL, isolation par établissement", "Sérieux, portable, auditable"],
            ["Identité", "JWT, rôles, révocation après changement de mot de passe", "Sessions maîtrisées"],
            ["Métiers sensibles", "PDF bulletins, Excel, imports CSV, e-mails", "Le quotidien réel de l’école, pas une démo"],
            ["Déploiement", "Cloud (web + API + base managée)", "Pas de serveur à installer dans chaque école"],
        ],
        [36, 68, 70],
    )
    pdf.bullets([
        "Isolation multi-école : une action ne peut pas modifier « l’autre établissement » par erreur de paramètre.",
        "Contrôle d’accès sur les fichiers (PDF, exports) : pas de lien public oublié.",
        "Limitation des tentatives de connexion.",
        "Mot de passe oublié par lien signé, ou reset en présentiel pour les comptes sans e-mail réel.",
        "Journal d’audit des actes sensibles (paiements, bulletins, utilisateurs).",
    ])
    pdf.body("En langage clair : Sulungukutu est déjà au niveau où l’on peut ouvrir une école pilote sans reconstruire le moteur. Le travail restant est de commercialisation, d’intégrations de paiement réelles, et d’accompagnement — pas de « refaire le logiciel ».")

    pdf.h1("11  ·  État d’avancement et feuille de route")
    pdf.h2("Ce qui est déjà là — sans maquillage")
    pdf.body("Nous assumons le stade : produit abouti, commercialisation à industrialiser. C’est un stade sain. Il évite de lever de l’argent pour découvrir, deux ans plus tard, que le métier école est plus complexe qu’un tableau de tâches.")
    pdf.table(
        ["Brique", "Statut"],
        [
            ["Espaces direction, enseignant, parent, élève, super-admin", "Livré"],
            ["Référentiel national Congo (niveaux, séries, matières, coefficients)", "Livré"],
            ["Notes, présences, bulletins PDF, publication", "Livré"],
            ["Mensualités, acomptes, modes de règlement, parcours Mobile Money (socle)", "Livré"],
            ["Messagerie, annonces, notifications temps réel", "Livré"],
            ["Onboarding école, imports, exports, audit", "Livré"],
            ["Identité, sécurité, multi-établissements", "Livré"],
            ["Premiers établissements payants en production", "À enclencher avec les partenaires de ce dossier"],
            ["Mobile Money opérateur (Airtel Money / MTN MoMo) en réel", "Prochaine étape prioritaire"],
            ["Application mobile native parents / enseignants", "Feuille de route"],
            ["SMS de masse", "Feuille de route"],
        ],
        [118, 56],
    )
    pdf.h2("Feuille de route")
    pdf.table(
        ["Période", "Priorité", "Résultat visé"],
        [
            ["0 – 6 mois", "Pilotes Brazzaville & Pointe-Noire, formation, preuves d’usage", "Études de cas : temps gagné, recouvrement, satisfaction parents"],
            ["6 – 12 mois", "Mobile Money production, SMS relances, offre annuelle rentrée", "Paiement à distance réel + renouvellements"],
            ["12 – 24 mois", "Réseau 50+ écoles, app mobile parents, second pays CEMAC", "Standard de fait au Congo privé urbain"],
            ["24 – 36 mois", "Modules RH / cantine, analytique avancée, groupes régionaux", "Plateforme d’exploitation scolaire, plus seulement un ENT"],
        ],
        [32, 72, 70],
    )

    pdf.h1("12  ·  Partenariats recherchés et prochaines étapes")
    pdf.h2("Ce que nous proposons, concrètement")
    pdf.body("Selon le profil du destinataire, plusieurs portes d’entrée existent. Elles ne s’excluent pas.")
    pdf.table(
        ["Vous êtes…", "Ce que nous vous proposons", "Ce que vous y gagnez"],
        [
            ["Promoteur ou réseau d’écoles", "Statut de pilote fondateur, tarif préférentiel, co-construction", "Un outil en avance, une voix sur la feuille de route, un avantage auprès des familles"],
            ["Association de promoteurs", "Partenariat de déploiement, démonstrations, offre réseau", "Un service concret pour vos membres, une modernisation visible"],
            ["Opérateur telecom / Mobile Money", "Intégration des paiements de scolarité dans Sulungukutu", "Un flux récurrent, ancré dans un usage quotidien des familles"],
            ["Institution financière / fintech", "Écoles plus lisibles, recouvrement plus clair (avec accord)", "Moins de risque sur les crédits rentrée, un canal d’écoles structurées"],
            ["Investisseur / business angel", "Entrée au capital ou financement d’amorçage go-to-market", "Un produit déjà construit, un marché concentré, un modèle SaaS + paiements"],
            ["Collaborateur clé", "Déploiement, commercial, succès client, partnerships", "Une entreprise naissante sur un besoin réel, pas une idée de salon"],
        ],
        [46, 64, 64],
    )
    pdf.h2("Emploi indicatif d’un financement d’amorçage")
    pdf.body("À titre de cadrage — le montant exact et les modalités se discutent — une enveloppe de l’ordre de 25 à 40 millions FCFA sur 12 à 18 mois permettrait d’industrialiser sans diluer l’attention :")
    pdf.table(
        ["Poste", "Part", "À quoi ça sert"],
        [
            ["Commercialisation et onboarding", "≈ 40 %", "Démos, déplacements Brazza / PNR, formation, succès des pilotes"],
            ["Paiements réels et SMS", "≈ 25 %", "Intégrations Airtel / MTN, passerelle SMS, tests de recouvrement"],
            ["Produit (app parents, durcissement)", "≈ 20 %", "Aller chercher les familles là où elles sont : le téléphone"],
            ["Infra, juridique, support", "≈ 15 %", "Hébergement, contrats, assistance rentrée"],
        ],
        [58, 28, 88],
    )
    pdf.h2("Les trois prochaines réunions qui comptent")
    pdf.bullets([
        "Une démonstration live de 45 minutes, sur l’école de démonstration, dans les cinq rôles.",
        "La visite d’un établissement candidat pilote : import des élèves, paramétrage, formation d’une matinée.",
        "Un protocole d’accord (pilote, distribution, ou discussion d’investissement) avec un calendrier de rentrée.",
    ])
    pdf.callout(
        "Invitation",
        "Si vous dirigez une école, un réseau, un opérateur ou un capital, Sulungukutu n’a pas besoin que vous « y croyiez » en l’air. Asseyez-vous devant le produit. Demandez le tableau des impayés, un bulletin, l’espace d’un parent. Puis décidez. C’est la seule séquence honnête — et c’est celle que nous proposons.",
    )

    pdf.h1("Contact")
    pdf.set_font(pdf.font_title, "B", 14)
    pdf.set_text_color(34, 31, 29)
    pdf.cell(0, 8, "Kassy Gloire Exaucé", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(pdf.font_body, "B", 11)
    pdf.set_text_color(79, 70, 229)
    pdf.cell(0, 6, "Fondateur — Sulungukutu", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(pdf.font_body, "", 10.5)
    pdf.set_text_color(28, 25, 23)
    pdf.cell(0, 6, "Plateforme de gestion scolaire pour l’Afrique centrale", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    pdf.muted("Document de présentation partenaires  ·  août 2026  ·  usage confidentiel. Ce document ne constitue pas une offre au public. Les tarifs, projections et scénarios de recouvrement sont fournis à titre illustratif pour éclairer une discussion. Toute collaboration fera l’objet d’un accord écrit.")

    pdf.output(str(PDF_PATH))


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_docx()
    print(f"DOCX  {DOCX_PATH}")
    build_pdf()
    print(f"PDF   {PDF_PATH}")
